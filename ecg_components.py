# =============================================================================
# 📈 Coleta e Visualização de ECG (ESP32 + AD8232) no Streamlit — versão enxuta
# -----------------------------------------------------------------------------
# - Lê linhas "ECG:12,34,56" ou "ECG: <adc>" da serial (Bluetooth/USB).
# - Converte ADC -> mV (apenas para exibição/CSV).
# - Mantém janela deslizante (10 s) para plot ao vivo.
# - Estima BPM por picos (SciPy se disponível; há fallback).
# - Gera DOCX com traçado, dados do paciente, resultado da IA e nota em letras miúdas.
# =============================================================================

from __future__ import annotations
import io
import time
from datetime import datetime
from typing import Optional

import numpy as np
import pandas as pd
import streamlit as st
import matplotlib
matplotlib.use("Agg")  # backend estável para apps sem GUI
import matplotlib.pyplot as plt


# Serial
try:
    import serial
    from serial.tools import list_ports
    HAS_SERIAL = True
except Exception:
    HAS_SERIAL = False

# SciPy opcional
try:
    from scipy.signal import find_peaks
    SCIPY_OK = True
except Exception:
    SCIPY_OK = False

# -----------------------
# Parâmetros centrais
# -----------------------
FS_DEFAULT     = 250.0   # Hz presumidos do firmware
PLOT_WINDOW_S  = 10      # janela do gráfico ao vivo (segundos)
CONCAT_DF      = True    # registra cada amostra no DataFrame

# ===================================
# 1) ESTADO DE SESSÃO
# ===================================
def init_session_state():
    """Inicializa chaves usadas ao longo do app."""
    ss = st.session_state
    if "raw_rows" not in ss:        ss.raw_rows = []          # buffers do RAW (lista de dicts)
    if "ser" not in ss:             ss.ser = None
    if "connected" not in ss:       ss.connected = False
    if "collecting" not in ss:      ss.collecting = False
    if "df" not in ss:
        ss.df = pd.DataFrame(
            {
                "tempo": pd.Series(dtype="string"),
                "adc":   pd.Series(dtype="float64"),
                "mv":    pd.Series(dtype="float64"),
            }
        )
    if "raw_buffer" not in ss:      ss.raw_buffer = []        # últimos N (mV) para plot
    if "bpm" not in ss:             ss.bpm = None
    if "fs" not in ss:              ss.fs = FS_DEFAULT
    if "selected_port" not in ss:   ss.selected_port = ""
    if "console_lines" not in ss:   ss.console_lines = []
    if "rxbuf" not in ss:           ss.rxbuf = b""
    if "last_render" not in ss:     ss.last_render = 0.0

# ===================================
# 2) SERIAL: PORTAS / CONEXÃO
# ===================================
def list_serial_ports():
    if not HAS_SERIAL:
        return []
    return [p.device for p in list_ports.comports()]

def connect_to_serial(port: str, baud: int):
    """Abre a serial e devolve (ser, msg_de_status) sem imprimir na tela."""
    if not HAS_SERIAL:
        return None, "Biblioteca pyserial ausente."
    try:
        ser = serial.Serial(port, baud, timeout=0, write_timeout=0)
        ser.reset_input_buffer()
        return ser, f"Conectado em {port}"
    except Exception as e:
        return None, f"Falha ao abrir {port}: {e}"

def disconnect_from_serial():
    try:
        if st.session_state.ser:
            st.session_state.ser.close()
            st.session_state.ser = None
        return "Desconectado com sucesso."
    except Exception as e:
        return f"Erro ao desconectar: {e}"

# ===================================
# 3) PARSING DE LINHAS
# ===================================
def parse_line(line: str):
    """
    Suporta:
      - 'ECG: <adc>'
      - 'ECG_raw:<adc>, ECG_smooth:<adc2>' (usa 'smooth' se existir)
    Retorna (valor_para_grafico, adc_raw, adc_smooth)
    Observação: LINHAS 'RAW:' são tratadas em _process_bt_buffer().
    """
    line = line.strip()
    if not line:
        return None, None, None

    if line.startswith("RAW:"):
        return None, None, None

    try:
        if line.startswith("ECG:"):
            val = float(line.split(":", 1)[1].strip())
            return val, val, None

        if "ECG_raw" in line and "ECG_smooth" in line:
            parts = line.split(",")
            raw = float(parts[0].split(":", 1)[1].strip())
            smooth = float(parts[1].split(":", 1)[1].strip())
            return smooth, raw, smooth
    except Exception:
        pass
    return None, None, None

# ===================================
# 4) CONVERSÃO & BPM
# ===================================
def adc_to_millivolts(adc: float, vref: float = 3.3, nbits: int = 12, gain: float = 1.0):
    if adc is None:
        return None
    mv = (adc / (2**nbits - 1)) * vref * 1000.0
    return mv / gain

def estimate_bpm(signal_mv, fs: float):
    sig = np.array(signal_mv, dtype=float)
    if len(sig) < int(2 * fs):
        return None

    x = (sig - sig.mean()) / (sig.std() + 1e-9)
    height = 0.6 * np.max(x)
    distance = int(0.25 * fs)

    if SCIPY_OK:
        peaks, _ = find_peaks(x, height=height, distance=distance)
    else:
        peaks = []
        last = -distance
        for i in range(1, len(x) - 1):
            if i - last < distance:
                continue
            if x[i] > height and x[i] > x[i-1] and x[i] > x[i+1]:
                peaks.append(i)
                last = i
        peaks = np.array(peaks, dtype=int)

    if len(peaks) < 2:
        return None

    rr = np.diff(peaks) / fs
    rr = rr[(rr > 0.3) & (rr < 2.0)]  # 30–200 BPM
    if len(rr) == 0:
        return None

    bpm = 60.0 / np.median(rr)
    return float(bpm) if 30 <= bpm <= 220 else None

# ===================================
# 5) LOOP PRINCIPAL DE COLETA
# ===================================
def start_collection(duration_s: float = 30.0, plot_ph=None, bpm_ph=None, console_ph=None):
    if not st.session_state.ser:
        st.error("Sem conexão serial.")
        return
    if st.session_state.collecting:
        st.warning("Coleta já em andamento.")
        return

    plot_ph = plot_ph or st.empty()
    bpm_ph  = bpm_ph or st.empty()
    st.session_state.collecting   = True
    st.session_state.rxbuf        = b""
    st.session_state.last_render  = 0.0
    start_time = time.time()
    # --- Inicializa variáveis de tempo e buffers ---
    st.session_state.df = pd.DataFrame(columns=["tempo","t_s","adc","mv"])
    st.session_state.raw_buffer = []
    st.session_state.raw_rows = []
    st.session_state.t0 = time.time()
    st.session_state.coleta_max_s = 35.0   # duração total
    st.session_state.descartar_s = 5.0     # segundos a descartar no início

    if "time_ph" in st.session_state:
        st.session_state.time_ph.caption("⏱️ Tempo decorrido: 0.0 s")

    msg_ph = st.empty()
    msg_ph.success(f"Coleta iniciada (tempo máximo: {duration_s:.0f} s).")
    msg_hide_at = time.time() + 3.0

    try:
        while st.session_state.collecting:
            n = getattr(st.session_state.ser, "in_waiting", 0)
            if n:
                chunk = st.session_state.ser.read(n)
                if chunk:
                    st.session_state.rxbuf += chunk
                    _process_bt_buffer()

            now = time.time()
            if now - st.session_state.last_render >= 0.05:
                _render(plot_ph, bpm_ph, console_ph)
                st.session_state.last_render = now

            if msg_ph is not None and time.time() >= msg_hide_at:
                msg_ph.empty()
                msg_ph = None

            elapsed = now - start_time
            if "time_ph" in st.session_state:
                st.session_state.time_ph.caption(f"⏱️ Tempo decorrido: {elapsed:5.1f} s")

            if now - start_time >= st.session_state.coleta_max_s:
                st.session_state.collecting = False
                st.success("Coleta concluída (35 s).")
                break


    except Exception as e:
        st.session_state.collecting = False
        st.error(f"Erro na leitura: {e}")

def _render(plot_placeholder, bpm_placeholder, console_placeholder=None):
    y = np.array(st.session_state.raw_buffer, dtype=float) if len(st.session_state.raw_buffer) >= 5 else np.array([])
    fs = st.session_state.fs
    x = np.arange(len(y)) / fs if len(y) > 0 else np.linspace(0, 10, 100)

    plt.style.use("dark_background")
    fig, ax = plt.subplots(figsize=(9, 3), facecolor="none")
    ax.set_facecolor("#FFFFFF")

    if len(y) > 0:
        ax.plot(x, y, linewidth=1.2, color="#DA4E3B")
    else:
        ax.plot(x, np.zeros_like(x), color="#66666683", linestyle="--", linewidth=0.8)
        ax.text(0.5, 0, "Aguardando amostras...", ha="center", va="center",
                color="#AAAAAA", fontsize=10)

    ax.set_ylim(-1000, 2000)
    ax.set_xlabel("t (s)", color="#CCCCCC")
    ax.set_ylabel("ECG (mV)", color="#CCCCCC")
    ax.tick_params(colors="#AAAAAA")
    ax.grid(True, color="#333333", alpha=0.5)
    fig.tight_layout()

    plot_placeholder.pyplot(fig)  # 1.43 não aceita 'width' aqui
    plt.close(fig)

    if st.session_state.bpm is not None:
        bpm_placeholder.metric("BPM (estimado)", f"{st.session_state.bpm:.0f}")
    else:
        bpm_placeholder.info("BPM: estimando...")

    if console_placeholder is not None and st.session_state.console_lines:
        console_placeholder.code("\n".join(st.session_state.console_lines[-10:]), language="text")

def render_live_plot_and_metrics(plot_ph, bpm_ph, console_ph=None):
    _render(plot_ph, bpm_ph, console_ph)

def stop_collection():
    if st.session_state.collecting:
        st.session_state.collecting = False
        st.success("Coleta parada.")
    else:
        st.warning("Nenhuma coleta ativa.")

# ===================================
# 6) UTILIDADES
# ===================================
def test_read(n: int = 10):
    if not st.session_state.ser:
        st.error("Sem conexão serial.")
        return
    lines = []
    st.info(f"Lendo {n} linhas...")
    for _ in range(n):
        line = st.session_state.ser.readline().decode("utf-8", errors="ignore").strip()
        lines.append(line if line else "<timeout>")
    st.code("\n".join(lines), language="text")

def download_csv_button():
    if st.session_state.df.empty:
        st.caption("Nenhum dado para exportar.")
        return
    buf = io.StringIO()
    st.session_state.df.to_csv(buf, index=False)
    st.download_button(
        "Baixar CSV",
        data=buf.getvalue(),
        file_name=f"ecg_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
        mime="text/csv"
    )

# ===================================
# 7) RX BUFFER
# ===================================
def _process_bt_buffer():
    buf = st.session_state.rxbuf
    if not buf:
        return

    lines = buf.split(b"\n")
    st.session_state.rxbuf = lines[-1]

    for rawline in lines[:-1]:
        line = rawline.decode("utf-8", "ignore").strip()
        if not line:
            continue

        st.session_state.console_lines.append(f"RAW_LINE: {repr(line)}")

        if line.startswith("RAW:"):
            payload = line.split(":", 1)[1]
            items = [s for s in payload.split(";") if s]
            for it in items:
                parts = [p.strip() for p in it.split(",")]
                if len(parts) >= 2:
                    try:
                        t_ms   = float(parts[0])
                        raw    = float(parts[1])
                        lo_p   = int(parts[2]) if len(parts) >= 3 else 0
                        lo_m   = int(parts[3]) if len(parts) >= 4 else 0
                        st.session_state.raw_rows.append(
                            {"t_ms": t_ms, "raw": raw, "lo_plus": lo_p, "lo_minus": lo_m}
                        )
                    except Exception:
                        pass
            continue

        if line.startswith("ECG:"):
            payload = line.split(":", 1)[1]
            parts = [p for p in payload.replace(";", ",").split(",") if p]
            for p in parts:
                try:
                    v = float(p)
                except ValueError:
                    continue
                _register_sample(v)
        else:
            val_used, raw, smooth = parse_line(line)
            if val_used is not None:
                _register_sample(val_used)

def _register_sample(adc_val: float):
    # tempo real desde o início da coleta
    now = time.time()
    t0 = st.session_state.get("t0", now)
    t_s = float(now - t0)  # tempo relativo em segundos
    ts = datetime.fromtimestamp(now).strftime("%d/%m/%Y %H:%M:%S.%f")[:-3]
    mv = adc_to_millivolts(adc_val)

    # grava no DataFrame
    if CONCAT_DF:
        row = {"tempo": ts, "t_s": t_s, "adc": float(adc_val), "mv": float(mv) if mv is not None else np.nan}
        df = st.session_state.df
        if df.empty:
            st.session_state.df = pd.DataFrame([row], columns=["tempo", "t_s", "adc", "mv"])\
                                     .astype({"tempo": "string", "t_s": "float64", "adc": "float64", "mv": "float64"})
        else:
            st.session_state.df.loc[len(df)] = row

    # janela de 10 s para o gráfico ao vivo
    st.session_state.raw_buffer.append(mv)
    max_points = int(st.session_state.fs * PLOT_WINDOW_S)
    if len(st.session_state.raw_buffer) > max_points:
        st.session_state.raw_buffer = st.session_state.raw_buffer[-max_points:]

    st.session_state.bpm = estimate_bpm(st.session_state.raw_buffer, st.session_state.fs)
    st.session_state.console_lines.append(f"ECG:{int(adc_val)}")
    if len(st.session_state.console_lines) > 200:
        st.session_state.console_lines = st.session_state.console_lines[-200:]


# =========================================
# 7.5) EXPORTAÇÃO NO FORMATO DO NOTEBOOK
# =========================================
def _cache_notebook_payload(df: pd.DataFrame, file_name: Optional[str] = None) -> None:
    st.session_state.notebook_df = df.copy()
    st.session_state.notebook_csv_bytes = df.to_csv(index=False).encode("utf-8")
    st.session_state.notebook_csv_name = file_name or f"entrada_notebook_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"

def get_notebook_df(clear: bool = False) -> pd.DataFrame:
    rows = st.session_state.get("raw_rows", [])
    if not rows:
        df = pd.DataFrame(columns=["t_ms", "value", "lo_plus", "lo_minus"])
    else:
        df_raw = pd.DataFrame(rows)
        if "raw" in df_raw.columns:
            df_raw = df_raw.rename(columns={"raw": "value"})
        keep = ["t_ms", "value", "lo_plus", "lo_minus"]
        for k in keep:
            if k not in df_raw.columns:
                df_raw[k] = 0
        df = (df_raw[keep]
              .drop_duplicates(subset=["t_ms"])
              .sort_values("t_ms")
              .reset_index(drop=True))
    _cache_notebook_payload(df)
    if clear:
        st.session_state.raw_rows = []
    return df

def save_notebook_csv(path: str, clear: bool = False) -> str:
    df = get_notebook_df(clear=clear)
    df.to_csv(path, index=False)
    st.session_state.notebook_csv_name = path
    return path

def get_cached_notebook_df() -> pd.DataFrame:
    if "notebook_df" in st.session_state:
        return st.session_state.notebook_df.copy()
    return pd.DataFrame(columns=["t_ms", "value", "lo_plus", "lo_minus"])

def get_cached_notebook_csv() -> tuple[bytes, str]:
    csv_bytes = st.session_state.get("notebook_csv_bytes", b"t_ms,value,lo_plus,lo_minus\n")
    file_name = st.session_state.get("notebook_csv_name", f"entrada_notebook_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv")
    return csv_bytes, file_name

# ===================================
# 8) EXPORTS (PNG/DOCX)
# ===================================
def make_ecg_plot_png(mode: str = "window"):
    if mode == "all" and not st.session_state.df.empty:
        df = st.session_state.df.copy()
        x = df["t_s"].to_numpy(dtype=float)
        y = df["mv"].to_numpy(dtype=float)

        # descarta primeiros 5 s e corta até 35 s
        start_cut = float(st.session_state.get("descartar_s", 5.0))
        dur = float(st.session_state.get("coleta_max_s", 35.0))
        keep = (x >= start_cut) & (x <= dur)
        x = x[keep] - start_cut
        y = y[keep]

        titulo = f"ECG – coleta de {dur:.0f}s (descartados {start_cut:.0f}s iniciais)"
    else:
        if len(st.session_state.raw_buffer) < 5:
            return None
        y = np.asarray(st.session_state.raw_buffer, dtype=float)
        fs = float(st.session_state.fs)
        x = np.arange(len(y)) / fs
        titulo = "ECG – últimos 10 s"

    fig, ax = plt.subplots(figsize=(9, 3))
    fig.patch.set_alpha(0.0)
    ax.set_facecolor("none")
    ax.plot(x, y, linewidth=2.0, color="#DA4E3B", antialiased=True)

    ax.set_xlabel("t (s)", color="black")
    ax.set_ylabel("ECG (mV)", color="black")
    ax.set_title(titulo, color="black")
    ax.grid(True, alpha=0.25, color="black")
    ax.tick_params(colors="black")
    for spine in ax.spines.values():
        spine.set_edgecolor("black")

    fig.tight_layout()
    bio = io.BytesIO()
    fig.savefig(bio, format="png", dpi=300, bbox_inches="tight", transparent=True)
    plt.close(fig)
    bio.seek(0)
    return bio.getvalue()


def build_docx_report(patient: dict, mode: str = "window"):
    """
    Gera DOCX com dados do paciente, traçado, resultado da IA e nota em letras miúdas.
    Usa st.session_state["last_ai_result"] se existir.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.oxml.ns import qn
    except Exception:
        st.error("Para exportar DOCX, instale:  pip install python-docx")
        return None

    doc = Document()
    doc.add_heading("Relatório de ECG com identificação de arritmia por IA", 0)
    doc.add_paragraph(f"Data/hora: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    # Dados do paciente
    p = patient or {}
    table = doc.add_table(rows=0, cols=2)
    def row(k, v):
        r = table.add_row().cells
        r[0].text = k
        r[1].text = v if (v is not None and str(v).strip() != "") else "—"

    row("Nome",        p.get("nome"))
    row("Idade",       str(p.get("idade")) if p.get("idade") is not None else None)
    row("Gênero",      p.get("genero"))
    row("Altura (cm)", str(p.get("altura_cm")) if p.get("altura_cm") else None)
    row("Peso (kg)",   f"{p.get('peso_kg'):.1f}" if p.get("peso_kg") else None)
    doc.add_paragraph("")

    # Resultado da IA (se disponível)
    result = st.session_state.get("last_ai_result", None)
    doc.add_heading("Resultado da Análise por IA", level=1)
    if result:
        label = result.get("label")
        prob = result.get("score", 0) * 100
        thr = result.get("threshold", 0) * 100
        if label == 1:
            txt = f"ARRITMIA DETECTADA (probabilidade {prob:.1f}% ≥ limiar {thr:.1f}%)"
        else:
            txt = f"SINAL NORMAL (probabilidade {prob:.1f}% < limiar {thr:.1f}%)"
        doc.add_paragraph(txt)
    else:
        doc.add_paragraph("Nenhum resultado disponível. Execute a classificação antes de gerar o relatório.")
    doc.add_paragraph("")

    # Traçado
    png = make_ecg_plot_png(mode=mode)
    if png:
        doc.add_paragraph("Traçado registrado:")
        doc.add_picture(io.BytesIO(png), width=Inches(6.0))
    else:
        doc.add_paragraph("Não há traçado disponível para inclusão.")

    # Nota em letras miúdas
    note = (
        "Este relatório é gerado automaticamente por um modelo de inteligência artificial em fase experimental. "
        "O resultado apresentado pode conter erros e não substitui avaliação médica ou exame oficial."
    )
    pnote = doc.add_paragraph()
    run = pnote.add_run(note)
    run.font.size = Pt(7)
    run.italic = True
    run.font.name = "Arial"
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()
